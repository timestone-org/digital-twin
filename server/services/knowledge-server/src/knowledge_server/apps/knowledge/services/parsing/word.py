"""Word 文档（.docx）：按文档序穿过段落与表格。

⚠ 段落与表格必须按**文档序**走：`Document.paragraphs` 与 `Document.tables`
是两份平行清单，各走各的会把全部表格堆到文末，而它们的标题路径于是取自最后
一个标题——第一章的表被引用成末章的。`iter_inner_content()` 才给真实先后。

⚠ 取**文本、结构与嵌入的图片**；不跑宏、不跟外部引用。图按它所在的那一段
定位（`§P3`），所以引用面上「这句话旁边配的是哪张图」是准的。

⚠ 三件事各有决定，见 ADR-0043：页眉页脚**每节收一次**（文件号与版本常常只
写在那儿，而它们在文件里本来就只存一份）；脚注尾注**够不着**（python-docx
1.2 没有公开面，钻 footnotes part 等于把库的内部结构写进这里）；批注**刻意
不收**——它是审阅过程的对话，收进来会让检索答出一份还没定稿的说法。

⚠ 解析是**纯 CPU 且阻塞**的，调用方必须扔进进程池。
"""

import re
from collections.abc import Sequence
from dataclasses import dataclass, field
from io import BytesIO
from typing import cast

from docx import Document
from docx.document import Document as DocxDocument
from docx.oxml.ns import qn
from docx.oxml.table import CT_Tbl
from docx.oxml.text.paragraph import CT_P
from docx.table import Table
from docx.text.paragraph import Paragraph

from knowledge_server.apps.knowledge.services.parsing.ports import (
    Block,
    Figure,
    Locator,
    ParsedDocument,
    RawItem,
)
from knowledge_server.apps.knowledge.services.parsing.structure import (
    cell_text,
    paired,
    path_of,
    pushed,
)

# 与纯文本那一路同一个上限，理由也同源
MAX_BLOCKS = 20_000
# 一份文档最多收几张图。⚠ 要有上限：图跟着正文一起进内存，一份塞了几千张
# 小图的文件会变成几千次对象存储写
MAX_FIGURES = 500
# 没有图注的图拿这句占位。⚠ 块的正文不许为空（`text_present` 那条 CHECK），
# 而一张没有图注的插图仍然值得在引用面上摆出来
NO_CAPTION = "（图，无图注）"
# 图注那一段的开头：「图 1」「表 2-1」「Figure 3」。⚠ 只认这几个词打头**且**
# 后面跟着编号的——认宽了会把图后面那一整段正文挂成图注，而错的图注比没有
# 图注更难发现
_CAPTION_HEAD = re.compile(
    r"^(图|表|附图|Figure|Fig\.?|Table)\s*[0-9０-９一二三四五六七八九十]"
)

_DOCX_SUFFIXES = (".docx",)
# 列表项的样式名前缀。⚠ 中英各一套：现场的模板多半是中文版 Word 存的
_LIST_STYLES = (
    "List Paragraph",
    "List Bullet",
    "List Number",
    "列表段落",
    "列表项目符号",
    "列表编号",
)


def _heading_level(style_name: str) -> int:
    """样式名 → 标题层级；不是标题给 0。

    ⚠ 认 `Heading N` 也认中文的「标题 N」：现场的模板多半是中文版 Word 存的，
    只认英文的话整份文档一个标题都解不出来，而那正是切块质量的主要来源。

    Args: style_name。
    """
    for prefix in ("Heading ", "标题 ", "标题"):
        if style_name.startswith(prefix):
            tail = style_name[len(prefix) :].strip()
            if tail.isdigit():
                return int(tail)
    return 0


def _style_name(paragraph: Paragraph) -> str:
    """这一段挂的样式名；没挂给空串。

    Args: paragraph。
    """
    style = paragraph.style
    return "" if style is None else (style.name or "")


def _list_level(style_name: str) -> int:
    """列表项的嵌套深度，从 1 起；不是列表项给 0。

    ⚠ 深度只认样式名的尾数（`List Bullet 2`）。真正记着层级的是
    `w:numPr/w:ilvl`，而 python-docx 没有公开面够得着它——套着「列表段落」
    样式的多级列表因此一律算第 1 层，宁可少说也不硬猜一个层级出来。

    Args: style_name。
    """
    if not style_name.startswith(_LIST_STYLES):
        return 0
    tail = style_name[-1]
    return int(tail) if tail.isdigit() and tail != "0" else 1


def _merged(pairs: list[tuple[str, int]]) -> list[str]:
    """把横向合并的单元格折成一格。

    ⚠ `row.cells` 按**网格列**给：一个跨 3 列的单元格会连着出现 3 次。照单
    全收的话，一行合并表头变成「甲 | 甲 | 甲」，而它会折进后面的每一行。

    Args: pairs（每格的文本与它横跨几列）。
    """
    out: list[str] = []
    skip = 0
    for text, span in pairs:
        if skip > 0:
            skip -= 1
            continue
        out.append(cell_text(text))
        skip = max(0, span - 1)
    return out


@dataclass
class _Walk:
    """扫一遍文档时攒着的那几样。"""

    made: list[Block] = field(default_factory=list[Block])
    # 标题栈：层级 + 文本
    stack: list[tuple[int, str]] = field(default_factory=list[tuple[int, str]])
    # 已经摊过几张表，用来给 locator 编号
    tables: int = 0
    # 收到的图，按出现序
    figures: list[Figure] = field(default_factory=list[Figure])
    # Word 的关系 id → 这一份产出里的 `ref`
    refs: dict[str, str] = field(default_factory=dict[str, str])


def _body_items(document: DocxDocument) -> list[CT_P | CT_Tbl]:
    """正文里的段落与表格元素，按文档序。

    ⚠ `inner_content_elements` 在 python-docx 1.2 里没有类型标注，`Any` 在
    这一处**一次收敛**（code-style-python）：往外带的话每个取值点都要再判
    一次，而漏判的那一处会在运行期才炸。

    Args: document。
    """
    raw = cast(
        list[object],
        document.element.body.inner_content_elements,  # pyright: ignore[reportUnknownMemberType]
    )
    return [one for one in raw if isinstance(one, CT_P | CT_Tbl)]


def _embedded(element: CT_P) -> list[str]:
    """这一段里挂着的图，按出现序给它们的关系 id。

    ⚠ 走 `a:blip/@r:embed` 而不是 `Document.inline_shapes`：后者是一份拍平的
    清单，取得到图却取不到它在正文的哪一段，而「图在哪一段」正是引用要指的
    那个位置。

    ⚠ `xpath` 回的是 `Any`，所以在这个函数里一次收敛成字符串，不往外带——与
    `_textbox_texts` 同源，python-docx 没有够得着图的公开面。

    Args: element。
    """
    found = cast(list[object], element.xpath(".//a:blip/@r:embed"))
    return [one for one in found if isinstance(one, str) and one]


def _caption_at(
    document: DocxDocument, elements: Sequence[CT_P | CT_Tbl], index: int
) -> str:
    """紧挨着的下一段如果是图注就拿来，不是给空串。

    ⚠ 只看**下一段**且只认「图 1」「表 2-1」这种打头的：Word 的惯例是把图注
    单独写成图下面的一段，而它在正文里就是一段普通文字。认宽了会把图后面那
    一整段正文挂成图注，而错的图注比没有图注更难发现。

    ⚠ 图自己那一段的文字**不当图注**：它已经是一个正文块了，再当图注会让同
    一句话在引用卡片上出现两遍。

    Args: document, elements, index。
    """
    nxt = elements[index + 1] if index + 1 < len(elements) else None
    if not isinstance(nxt, CT_P):
        return ""
    text = Paragraph(nxt, document).text.strip()
    return text if _CAPTION_HEAD.match(text) else ""


def _figure_into(
    walk: _Walk, document: DocxDocument, rid: str, caption: str
) -> None:
    """一张图：收字节、去重，并出一个 `figure` 块指向它。

    ⚠ 同一个关系 id 在文里出现两次只收一份字节，两个块指同一个 `ref`：Word
    复用同一张图是常事，而重复上传的那一份到了对象存储那一步本来也会被内容
    哈希合掉。

    ⚠ 关系拿不到、字节是空的、或者根本不是图片的一律**跳过而不是整份失败**：
    一张坏图不该让一份两百页的手册摄不进来。

    Args: walk, document, rid, caption。
    """
    ref = walk.refs.get(rid, "")
    if not ref:
        if len(walk.figures) >= MAX_FIGURES:
            return
        part = document.part.related_parts.get(rid)
        blob = bytes(part.blob) if part is not None else b""
        media = str(part.content_type) if part is not None else ""
        if not blob or not media.startswith("image/"):
            return
        ref = f"docx-{len(walk.figures) + 1}"
        walk.refs[rid] = ref
        walk.figures.append(
            Figure(ref=ref, content=blob, media_type=media, caption=caption)
        )
    walk.made.append(
        Block(
            kind="figure",
            text=caption or NO_CAPTION,
            locator=Locator(path=path_of(walk.stack)),
            figure_ref=ref,
        )
    )


def _paragraph_into(
    walk: _Walk,
    document: DocxDocument,
    element: CT_P,
    caption: str,
) -> None:
    """一个段落成一块：标题压栈，列表项带上嵌套深度；段里挂的图另出块。

    ⚠ 先出正文块再出图块：图注多半写在图**后面**那一段，而块序就是引用面上
    那几处的先后。

    Args: walk, document, element, caption。
    """
    paragraph = Paragraph(element, document)
    _text_into(walk, paragraph)
    for rid in _embedded(element):
        _figure_into(walk, document, rid, caption)


def _text_into(walk: _Walk, paragraph: Paragraph) -> None:
    """段落的文字成一块；空段落不收。

    Args: walk, paragraph。
    """
    text = paragraph.text.strip()
    if not text:
        return
    name = _style_name(paragraph)
    level = _heading_level(name)
    if level > 0:
        walk.stack = pushed(walk.stack, level, text)
        walk.made.append(
            Block(
                kind="heading",
                text=text,
                level=level,
                locator=Locator(path=path_of(walk.stack)),
            )
        )
        return
    depth = _list_level(name)
    walk.made.append(
        Block(
            kind="list_item" if depth > 0 else "paragraph",
            text=text,
            level=depth,
            locator=Locator(path=path_of(walk.stack)),
        )
    )


def _table_into(walk: _Walk, table: Table) -> None:
    """一张表摊成行块，表头折进后面的每一行。

    Args: walk, table。
    """
    walk.tables += 1
    path = path_of(walk.stack)
    header: list[str] = []
    for index, row in enumerate(table.rows, start=1):
        if len(walk.made) >= MAX_BLOCKS:
            return
        cells = _merged([(one.text, one.grid_span) for one in row.cells])
        if not any(cells):
            continue
        walk.made.append(
            Block(
                kind="table_row",
                text=paired(header, cells) if header else " | ".join(cells),
                locator=Locator(
                    sheet=f"表 {walk.tables}", row=index, path=path
                ),
            )
        )
        header = header or cells


def _part_text(paragraphs: list[Paragraph]) -> str:
    """页眉或页脚里的几行拼成一句。

    Args: paragraphs。
    """
    return "\n".join(one.text.strip() for one in paragraphs if one.text.strip())


def _section_blocks(document: DocxDocument) -> list[Block]:
    """每一节的页眉页脚各收一次。

    ⚠ 只收一次而不是每页一次：它们在文件里本来就只存一份，Word 的重复是渲染
    时的事。按页收的话，一份 80 页的手册会多出 160 个几乎相同的块，把正文挤出
    召回名单。

    ⚠ 跳过「同前一节」的：那几节的页眉就是上一节那一份，收了就是重复；没定义
    页眉的文档也走这一支，于是不会多出一堆空块。

    Args: document。
    """
    made: list[Block] = []
    for index, section in enumerate(document.sections, start=1):
        for label, part in (
            ("页眉", section.header),
            ("页脚", section.footer),
        ):
            if part.is_linked_to_previous:
                continue
            text = _part_text(list(part.paragraphs))
            if text:
                made.append(
                    Block(
                        kind="caption",
                        text=text,
                        locator=Locator(sheet=label, row=index),
                    )
                )
    return made


def _textbox_texts(document: DocxDocument) -> list[str]:
    """文档里所有文本框的文字，按出现序去重。

    ⚠ 必须去重：Word 把同一个文本框在 `mc:Choice` 与 `mc:Fallback` 里各存
    一份，照单全收会让每一句都出现两遍。

    ⚠ `xpath` 回的是 `Any`，所以在这个函数里一次收敛成字符串，不往外带：
    python-docx 没有文本框的公开面，这是唯一够得着它的路。

    Args: document。
    """
    seen: list[str] = []
    for box in document.element.xpath(".//w:txbxContent"):
        for para in box.iter(qn("w:p")):
            text = "".join(
                str(node.text or "") for node in para.iter(qn("w:t"))
            ).strip()
            if text and text not in seen:
                seen.append(text)
    return seen


def _textbox_blocks(document: DocxDocument, made_so_far: int) -> list[Block]:
    """文本框里的文字，收在正文之后。

    ⚠ 现场的 Word 模板常把限值与注意事项放在文本框里，而文本框不在段落清单
    里——不捞的话那些内容一个字都进不来。

    ⚠ 收在末尾且**不带标题路径**：文本框是浮动对象，python-docx 没有公开面能
    把它定位到正文的哪一段，硬猜一个路径会让引用指错地方。

    Args: document, made_so_far。
    """
    made: list[Block] = []
    for index, text in enumerate(_textbox_texts(document), start=1):
        if made_so_far + len(made) >= MAX_BLOCKS:
            break
        made.append(
            Block(
                kind="paragraph",
                text=text,
                locator=Locator(sheet="文本框", row=index),
            )
        )
    return made


@dataclass(frozen=True)
class DocxParser:
    """Word 文档：按文档序走段落与表格，另收页眉页脚与文本框。"""

    name: str = "docx"
    suffixes: tuple[str, ...] = _DOCX_SUFFIXES
    media_types: tuple[str, ...] = (
        "application/vnd.openxmlformats-officedocument"
        ".wordprocessingml.document",
    )

    def parse(self, raw: RawItem) -> ParsedDocument:
        """页眉页脚 → 正文（段落、表格与插图按文档序）→ 文本框。

        ⚠ 走的是 `body.inner_content_elements` 而不是 `iter_inner_content()`：
        两者给的是同一串东西同一个次序，但前者连 XML 元素一起给——图挂在
        `w:drawing` 上，只拿 `Paragraph` 够不着它。

        Args: raw。
        """
        document = Document(BytesIO(raw.content))
        walk = _Walk()
        walk.made.extend(_section_blocks(document))
        elements = _body_items(document)
        for index, element in enumerate(elements):
            if len(walk.made) >= MAX_BLOCKS:
                break
            if isinstance(element, CT_P):
                _paragraph_into(
                    walk,
                    document,
                    element,
                    _caption_at(document, elements, index),
                )
            else:
                _table_into(walk, Table(element, document))
        walk.made.extend(_textbox_blocks(document, len(walk.made)))
        made = walk.made[:MAX_BLOCKS]
        return ParsedDocument(
            title=raw.filename,
            blocks=tuple(made),
            is_truncated=len(made) >= MAX_BLOCKS,
            figures=tuple(walk.figures),
        )
