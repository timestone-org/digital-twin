"""Word 解析。用 python-docx 现造样件，不放二进制夹具进仓。

⚠ 现造而不是放夹具文件：一份 .docx 是一个 zip，改一个字都看不出 diff，
而「夹具是怎么来的」半年后没人说得清。现造的代价是每条用例慢几毫秒。

⚠ 文本框那一份必须手拼 XML：python-docx 造不出文本框，而现场的 Word 模板
常把关键信息放在里面。
"""

from io import BytesIO

from docx import Document
from docx.document import Document as DocxDocument
from docx.enum.style import WD_STYLE_TYPE
from docx.oxml.ns import nsdecls
from docx.oxml.parser import parse_xml
from docx.table import Table

from knowledge_server.apps.knowledge.services.parsing import RawItem
from knowledge_server.apps.knowledge.services.parsing.word import (
    MAX_BLOCKS,
    DocxParser,
)

# 一个带 mc:AlternateContent 的文本框：Word 把同一段文字在 Choice 与 Fallback
# 里各存一份，两份都收就会重复
_TEXTBOX_RUN = (
    '<w:r %s xmlns:mc="http://schemas.openxmlformats.org/'
    'markup-compatibility/2006" xmlns:wps="http://schemas.microsoft.com/'
    'office/word/2010/wordprocessingShape" '
    'xmlns:v="urn:schemas-microsoft-com:vml">'
    "<mc:AlternateContent>"
    '<mc:Choice Requires="wps"><w:drawing><wps:txbx><w:txbxContent>'
    "<w:p><w:r><w:t>%s</w:t></w:r></w:p>"
    "</w:txbxContent></wps:txbx></w:drawing></mc:Choice>"
    "<mc:Fallback><w:pict><v:shape><v:textbox><w:txbxContent>"
    "<w:p><w:r><w:t>%s</w:t></w:r></w:p>"
    "</w:txbxContent></v:textbox></v:shape></w:pict></mc:Fallback>"
    "</mc:AlternateContent></w:r>"
)


def _raw(name: str, body: bytes) -> RawItem:
    return RawItem(filename=name, media_type="", content=body)


def _saved(document: DocxDocument) -> bytes:
    buffer = BytesIO()
    document.save(buffer)
    return buffer.getvalue()


def _parsed(document: DocxDocument) -> list[tuple[str, str]]:
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    return [(one.kind, one.text) for one in made.blocks]


def test_docx_headings_build_a_path() -> None:
    document = Document()
    document.add_paragraph("第一章", style="Heading 1")
    document.add_paragraph("1.1 冷却水", style="Heading 2")
    document.add_paragraph("出口温度不得高于 65 ℃")
    made = DocxParser().parse(_raw("手册.docx", _saved(document)))
    body = made.blocks[-1]
    assert body.text == "出口温度不得高于 65 ℃"
    assert body.locator.path == ("第一章", "1.1 冷却水")


def test_chinese_style_names_are_headings_too() -> None:
    """⚠ 现场的模板多半是中文版 Word 存的，只认 `Heading N` 的话整份文档
    一个标题都解不出来——而那正是切块质量的主要来源。"""
    document = Document()
    document.styles.add_style("标题 1", WD_STYLE_TYPE.PARAGRAPH).base_style = (
        document.styles["Heading 1"]
    )
    document.add_paragraph("第一章", style="标题 1")
    document.add_paragraph("正文")
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    assert made.blocks[0].kind == "heading"
    assert made.blocks[-1].locator.path == ("第一章",)


def test_empty_paragraphs_are_dropped() -> None:
    document = Document()
    for text in ("甲", "   ", "乙"):
        document.add_paragraph(text)
    assert _parsed(document) == [("paragraph", "甲"), ("paragraph", "乙")]


def _with_table(document: DocxDocument) -> Table:
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "点位"
    table.cell(0, 1).text = "上限"
    table.cell(1, 0).text = "出口温度"
    table.cell(1, 1).text = "65"
    return table


def test_a_table_keeps_its_place_in_the_document() -> None:
    """⚠ `paragraphs` 与 `tables` 是两份平行清单，各走各的会把全部表格堆到
    文末——于是「表后正文」跑到了表格前面，引用的先后全反了。"""
    document = Document()
    document.add_paragraph("表前正文")
    _with_table(document)
    document.add_paragraph("表后正文")
    assert [one[0] for one in _parsed(document)] == [
        "paragraph",
        "table_row",
        "table_row",
        "paragraph",
    ]
    assert _parsed(document)[-1] == ("paragraph", "表后正文")


def test_table_rows_inherit_the_heading_above_them() -> None:
    """⚠ 表格堆到文末的话，它们的标题路径取自**最后一个**标题——于是第一章
    的表被引用成第二章的。"""
    document = Document()
    document.add_paragraph("第一章", style="Heading 1")
    _with_table(document)
    document.add_paragraph("第二章", style="Heading 1")
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    rows = [one for one in made.blocks if one.kind == "table_row"]
    assert rows[0].locator.path == ("第一章",)


def test_the_table_header_is_folded_into_every_row() -> None:
    """⚠ 与工作簿那一路同一条规矩：只存 `出口温度 | 65` 的话，检索到这一行
    读不出 65 是什么——列名在表头那一行，而那一行是另一个块。"""
    document = Document()
    _with_table(document)
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    rows = [one for one in made.blocks if one.kind == "table_row"]
    assert rows[0].text == "点位 | 上限"
    assert rows[1].text == "点位=出口温度 | 上限=65"
    assert (rows[1].locator.sheet, rows[1].locator.row) == ("表 1", 2)


def test_merged_cells_are_not_repeated() -> None:
    """⚠ `row.cells` 按网格列给：一个跨 3 列的表头会连着出现 3 次，照单全收
    的话一行合并表头变成「甲 | 甲 | 甲」，而它会折进后面每一行。"""
    document = Document()
    table = document.add_table(rows=2, cols=3)
    table.cell(0, 0).merge(table.cell(0, 1))
    table.cell(0, 0).text = "冷却水"
    table.cell(0, 2).text = "备注"
    table.cell(1, 0).text = "甲"
    table.cell(1, 1).text = "乙"
    table.cell(1, 2).text = "丙"
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    rows = [one for one in made.blocks if one.kind == "table_row"]
    assert rows[0].text == "冷却水 | 备注"


def test_blank_table_rows_are_dropped() -> None:
    document = Document()
    table = document.add_table(rows=2, cols=2)
    table.cell(0, 0).text = "列"
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    assert len([one for one in made.blocks if one.kind == "table_row"]) == 1


def test_two_tables_get_distinct_locators() -> None:
    document = Document()
    _with_table(document)
    document.add_paragraph("中间")
    _with_table(document)
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    sheets = {
        one.locator.sheet for one in made.blocks if one.kind == "table_row"
    }
    assert sheets == {"表 1", "表 2"}


def test_list_paragraphs_keep_their_own_kind() -> None:
    """⚠ 列表项与普通段落混成一种的话，切块层看不出这几行是并列的几条。"""
    document = Document()
    document.add_paragraph("甲", style="List Bullet")
    document.add_paragraph("乙", style="List Number")
    document.add_paragraph("普通段落")
    assert _parsed(document) == [
        ("list_item", "甲"),
        ("list_item", "乙"),
        ("paragraph", "普通段落"),
    ]


def test_a_nested_list_style_reports_its_depth() -> None:
    document = Document()
    document.add_paragraph("甲", style="List Bullet")
    document.add_paragraph("甲一", style="List Bullet 2")
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    assert [one.level for one in made.blocks] == [1, 2]


def test_hyperlink_text_stays_in_the_paragraph() -> None:
    """⚠ 超链接的文字在 `w:hyperlink` 里而不在 `w:r` 里。这一条钉的是
    python-docx ≥ 1.1 把它算进 `Paragraph.text`——降级到更老的版本会让整段
    只剩链接前面那几个字，而没有任何一处报错。"""
    document = Document()
    paragraph = document.add_paragraph("详见 ")
    link = parse_xml(
        f"<w:hyperlink {nsdecls('w')}><w:r><w:t>运行手册</w:t></w:r>"
        "</w:hyperlink>"
    )
    paragraph._p.append(link)
    assert _parsed(document) == [("paragraph", "详见 运行手册")]


def test_textbox_text_is_collected_exactly_once() -> None:
    """⚠ 现场的 Word 模板常把限值与注意事项放在文本框里，而文本框不在段落
    清单里——不捞的话那些内容一个字都进不来。

    ⚠ 而 Word 把同一个文本框在 `mc:Choice` 与 `mc:Fallback` 里各存一份，
    照单全收会让每一句都出现两遍。"""
    document = Document()
    paragraph = document.add_paragraph("锚点段落")
    text = "出口温度上限 65 ℃"
    paragraph._p.append(parse_xml(_TEXTBOX_RUN % (nsdecls("w"), text, text)))
    made = _parsed(document)
    assert made.count(("paragraph", text)) == 1


def test_the_header_and_footer_are_collected_once_per_section() -> None:
    """⚠ 文件号与版本常常只写在页眉里，不收就再也检索不到；而它们在文件里
    本来就只存一份，收多了会在检索里制造几十份几乎相同的块。"""
    document = Document()
    section = document.sections[0]
    section.header.paragraphs[0].text = "文件号 Q-01"
    section.footer.paragraphs[0].text = "受控文件"
    document.add_paragraph("正文")
    made = _parsed(document)
    assert made.count(("caption", "文件号 Q-01")) == 1
    assert made.count(("caption", "受控文件")) == 1


def test_a_document_without_a_header_gets_no_caption() -> None:
    """⚠ 没定义页眉时 python-docx 给的是「继承上一节」的空壳，收进来就是一堆
    空块。"""
    document = Document()
    document.add_paragraph("正文")
    assert _parsed(document) == [("paragraph", "正文")]


def test_comments_are_deliberately_left_out() -> None:
    """⚠ 批注是审阅过程的对话（「这里要改」「同意」），收进来会让检索答出一份
    还没定稿的说法。python-docx 1.2 够得着它，所以这是选择不是限制。"""
    document = Document()
    paragraph = document.add_paragraph("正文")
    document.add_comment(paragraph.runs, "这一条要改成 70", author="审阅人")
    assert _parsed(document) == [("paragraph", "正文")]


def test_truncation_is_reported() -> None:
    """⚠ 悄悄截断的话，后面会把「我看到的就是全部」当成事实。"""
    document = Document()
    for index in range(MAX_BLOCKS + 5):
        document.add_paragraph(f"第 {index} 行")
    made = DocxParser().parse(_raw("a.docx", _saved(document)))
    assert made.is_truncated is True
    assert len(made.blocks) == MAX_BLOCKS
